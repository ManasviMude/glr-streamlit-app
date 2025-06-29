import fitz  # PyMuPDF
import json
import requests
import os
from io import BytesIO
from docx import Document
import streamlit as st

# === SECURELY LOAD API KEY ===
OPENROUTER_API_KEY = st.secrets.get("OPENROUTER_API_KEY") or os.getenv("OPENROUTER_API_KEY")
if not OPENROUTER_API_KEY:
    st.error("❌ OpenRouter API key not found. Add it in Streamlit > Settings > Secrets.")
    st.stop()

# Optional: required if your key is domain-locked (free keys)
HTTP_REFERER = "https://your-app-name.streamlit.app"  # Replace with your deployed URL on Streamlit Cloud

# === CLEAN TEXT TO REMOVE NON-UTF8 CHARACTERS ===
def clean_text(text):
    return text.encode("utf-8", "ignore").decode("utf-8")

# === EXTRACT TEXT FROM PDF ===
def extract_pdf_text(uploaded_pdfs):
    combined_text = ""
    for file in uploaded_pdfs:
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            for page in doc:
                combined_text += page.get_text()
    return clean_text(combined_text)

# === EXTRACT PLACEHOLDERS FROM DOCX TEMPLATE ===
def extract_placeholders(docx_file):
    doc = Document(docx_file)
    placeholders = set()
    for para in doc.paragraphs:
        for word in para.text.split():
            if word.startswith("[") and word.endswith("]"):
                placeholders.add(word.strip("[]"))
    return list(placeholders)

# === CALL LLM TO FILL PLACEHOLDER VALUES ===
def call_llm(pdf_text, placeholders):
    prompt = f"""
You are an insurance claim assistant. Extract values for the following placeholders:

{placeholders}

PDF Text:
\"\"\"
{pdf_text[:6000]}
\"\"\"

Return ONLY valid JSON. Example:
{{
  "DATE_LOSS": "2024-11-13",
  "INSURED_NAME": "Richard Daly",
  "MORTGAGE_CO": "Alacrity Mortgage",
  ...
}}
"""

    headers = {
        "Authorization": f"Bearer {OPENROUTER_API_KEY}",
        "Content-Type": "application/json",
        "HTTP-Referer": HTTP_REFERER  # Optional but good if your key is domain-locked
    }

    body = {
        "model": "mistralai/mixtral-8x7b",
        "messages": [{"role": "user", "content": prompt}]
    }

    try:
        body_json = json.dumps(body, ensure_ascii=False).encode("utf-8")  # ✅ proper encoding
        res = requests.post(
            "https://openrouter.ai/api/v1/chat/completions",
            headers=headers,
            data=body_json
        )
        res.raise_for_status()
        data = res.json()
        raw_output = data["choices"][0]["message"]["content"]
        return json.loads(raw_output)
    except requests.exceptions.HTTPError as e:
        st.error(f"HTTP {res.status_code} error: {res.text}")
    except Exception as e:
        st.error(f"LLM call failed: {e}")
    return {}

# === FALLBACK MOCK DATA ===
def mock_data():
    return {
        "DATE_LOSS": "2024-11-13",
        "INSURED_NAME": "Richard Daly",
        "MORTGAGE_CO": "Alacrity",
        "INSURED_H_STREET": "123 Storm Ln",
        "INSURED_H_CITY": "San Antonio",
        "INSURED_H_STATE": "TX",
        "INSURED_H_ZIP": "78265",
        "DATE_INSPECTED": "2024-11-14",
        "TOL_CODE": "wind",
        "DATE_RECEIVED": "2024-11-15",
        "MORTGAGEE": "Alacrity"
    }

# === FILL THE DOCX TEMPLATE WITH VALUES ===
def fill_template(docx_file, field_values):
    doc = Document(docx_file)
    for para in doc.paragraphs:
        for key, val in field_values.items():
            if f"[{key}]" in para.text:
                para.text = para.text.replace(f"[{key}]", val)
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# === STREAMLIT WEB APP ===
st.set_page_config(page_title="GLR Report Filler", page_icon="📝")
st.title("📄 GLR Pipeline - Auto-Fill Insurance Template")

template_file = st.file_uploader("Upload Template (.docx)", type=["docx"])
pdf_files = st.file_uploader("Upload Photo Reports (.pdf)", type=["pdf"], accept_multiple_files=True)

if st.button("Process"):
    if not template_file or not pdf_files:
        st.error("Please upload both a DOCX template and at least one PDF report.")
    else:
        with st.spinner("🔍 Extracting text from PDFs..."):
            pdf_text = extract_pdf_text(pdf_files)

        with st.spinner("🔎 Detecting placeholders from template..."):
            placeholders = extract_placeholders(template_file)

        with st.spinner("🤖 Calling LLM..."):
            field_values = call_llm(pdf_text, placeholders)
            if not field_values:
                st.warning("⚠️ LLM failed. Using mock data instead.")
                field_values = mock_data()

        st.success("✅ Fields extracted and matched!")

        with st.spinner("📝 Filling in template..."):
            filled_doc = fill_template(template_file, field_values)

        st.download_button(
            label="📥 Download Filled Report",
            data=filled_doc,
            file_name="filled_report.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
